import discord
from discord.ext import commands
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os
import re


class MarkdownConverter(commands.Cog):
    def __init__(self, bot):
        self.bot = bot

    @commands.command(name="convert_md")
    async def convert_md(self, ctx):
        if not ctx.message.attachments:
            await ctx.send("Please attach an .md file to convert.")
            return

        for attachment in ctx.message.attachments:
            if not attachment.filename.endswith(".md"):
                await ctx.send(f"{attachment.filename} is not a Markdown file.")
                continue

            # Download the file
            file_path = f"./{attachment.filename}"
            await attachment.save(file_path)

            # Convert Markdown to Word
            try:
                with open(file_path, "r", encoding="utf-8") as md_file:
                    content = md_file.read()

                document = Document()

                # Process the content and add to the Word document with formatting
                self.process_markdown(content, document)

                word_file_path = file_path.replace(".md", ".docx")
                document.save(word_file_path)

                # Send the Word file back
                await ctx.send(file=discord.File(word_file_path))

                # Clean up files
                os.remove(file_path)
                os.remove(word_file_path)
            except Exception as e:
                await ctx.send(f"An error occurred: {e}")

    def process_markdown(self, content, document):
        # Split content into lines to process individually
        lines = content.splitlines()

        for line in lines:
            # Handle Headings (Markdown syntax: # for H1, ## for H2, etc.)
            if line.startswith("#"):
                level = line.count("#")
                heading_text = line.lstrip("#").strip()
                self.add_heading(heading_text, level, document)
            # Handle Lists (Markdown syntax: - or * for unordered lists, numbers for ordered lists)
            elif line.startswith("- ") or line.startswith("* "):
                document.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r"^\d+\.", line):  # Ordered list
                document.add_paragraph(line[3:], style='List Number')
            # Handle code blocks (Markdown syntax: ```code block```)
            elif line.startswith("```"):
                code_block = self.extract_code_block(line, lines)
                document.add_paragraph(code_block, style='Normal')
            # Handle inline code (Markdown syntax: `code`)
            elif "`" in line:
                line = self.convert_inline_code(line)
                document.add_paragraph(line, style='Normal')
            # Handle bold and italic
            elif "**" in line or "__" in line:
                line = self.convert_bold_italic(line)
                document.add_paragraph(line, style='Normal')
            elif "*" in line or "_" in line:
                line = self.convert_bold_italic(line)
                document.add_paragraph(line, style='Normal')
            # Handle tables (Markdown syntax: | col1 | col2 |)
            elif "|" in line:
                if line.strip().startswith("|"):
                    self.add_table(line, lines, document)
            # Handle links (Markdown syntax: [text](url))
            elif "[" in line and "]" in line and "(" in line and ")" in line:
                line = self.convert_links(line)
                document.add_paragraph(line)
            # Regular paragraphs
            else:
                document.add_paragraph(line)

    def add_heading(self, text, level, document):
        """ Add a heading to the document based on the level. """
        if level == 1:
            document.add_heading(text, level=1)
        elif level == 2:
            document.add_heading(text, level=2)
        elif level == 3:
            document.add_heading(text, level=3)
        else:
            document.add_paragraph(text, style='Heading 4')

    def convert_bold_italic(self, line):
        """ Convert bold and italic Markdown syntax to plain text with formatting. """
        # Bold: Replace **bold** or __bold__ with Word formatting
        line = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", line)
        line = re.sub(r"__(.*?)__", r"<b>\1</b>", line)
        # Italic: Replace *italic* or _italic_ with Word formatting
        line = re.sub(r"\*(.*?)\*", r"<i>\1</i>", line)
        line = re.sub(r"_(.*?)_", r"<i>\1</i>", line)
        return line

    def convert_inline_code(self, line):
        """ Convert inline code (e.g. `code`) to Word formatted text. """
        line = re.sub(r"`(.*?)`", r"<code>\1</code>", line)
        return line

    def extract_code_block(self, current_line, lines):
        """ Extract multi-line code block from Markdown. """
        code_block = current_line.strip("```")  # Remove the code block indicators
        for line in lines:
            if line.startswith("```"):
                break
            code_block += "\n" + line.strip()  # Append code lines
        return code_block

    def add_table(self, header_line, lines, document):
        """ Convert Markdown tables to Word tables. """
        # Split the header line and the separator (---)
        headers = header_line.split("|")
        headers = [header.strip() for header in headers if header.strip()]

        # Create the table
        table = document.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header

        # Add rows to the table
        for line in lines:
            if "|" not in line or line.startswith("|") == False:
                continue
            row_data = line.split("|")
            row_data = [cell.strip() for cell in row_data if cell.strip()]
            row_cells = table.add_row().cells
            for i, data in enumerate(row_data):
                row_cells[i].text = data

    def convert_links(self, line):
        """ Convert Markdown links [text](url) to clickable links in Word. """
        # Example: [Google](https://www.google.com)
        match = re.search(r"\[([^\]]+)\]\(([^)]+)\)", line)
        if match:
            text = match.group(1)
            url = match.group(2)
            return f'{text} ({url})'
        return line


async def setup(bot):
    await bot.add_cog(MarkdownConverter(bot))
