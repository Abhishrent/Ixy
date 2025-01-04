import zipfile
import os
import tempfile
import re
import io
import discord
from discord.ext import commands
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
import asyncio

class LabReportGenerator(commands.Cog):
    def __init__(self, bot):
        self.bot = bot

    def extract_number_from_filename(self, filename):
        """Extract the number prefix from filename."""
        match = re.match(r'^(\d+)[.\s_-]', filename)
        return int(match.group(1)) if match else None

    def parse_questions(self, content):
        """Parse questions and their associated theory from the text content."""
        questions = {}
        try:
            # Split content into lines and process
            lines = content.split('\n')
            current_question = None
            current_theory = []

            for line in lines:
                line = line.strip()
                
                # Check for new question
                question_match = re.match(r'^(\d+)[.\s]+(.+)$', line)
                if question_match:
                    # If we were processing a previous question, save it
                    if current_question is not None:
                        questions[current_question]['theory'] = ' '.join(current_theory).strip()
                        current_theory = []

                    # Start new question
                    num, question = question_match.groups()
                    current_question = int(num)
                    questions[current_question] = {
                        'text': question.strip(),
                        'theory': ''
                    }
                    continue

                # Check for theory section
                if current_question and line.lower().startswith('theory:'):
                    # Skip the 'Theory:' line itself
                    continue

                # Collect theory lines
                if current_question and line:
                    current_theory.append(line)

            # Save theory for the last question
            if current_question is not None and current_theory:
                questions[current_question]['theory'] = ' '.join(current_theory).strip()

        except Exception as e:
            raise ValueError(f"Error parsing questions: {e}")
        
        return questions

    def extract_zip(self, zip_bytes):
        """Extract ZIP file contents and return a dictionary of file paths."""
        temp_dir = tempfile.TemporaryDirectory()
        extracted_files = {}
        try:
            with zipfile.ZipFile(io.BytesIO(zip_bytes)) as zip_file:
                zip_file.extractall(temp_dir.name)
                for root, _, files in os.walk(temp_dir.name):
                    for file in files:
                        file_path = os.path.join(root, file)
                        extracted_files[file] = file_path
        except zipfile.BadZipFile:
            raise ValueError("The uploaded file is not a valid ZIP file.")
        return extracted_files, temp_dir

    async def update_progress_embed(self, message, current_step, total_steps, details=None):
        """Update the progress embed with current step and optional details."""
        # Calculate progress percentage
        progress_percentage = int((current_step / total_steps) * 100)
        
        # Create progress bar
        progress_bar = self.create_progress_bar(progress_percentage)
        
        # Prepare embed description
        description = f"{progress_bar} {progress_percentage}%\n"
        if details:
            description += f"**Current Task:** {details}"
        
        embed = discord.Embed(
            title="Lab Report Generation Progress",
            description=description,
            color=discord.Color.blue()
        )
        
        # Add steps as fields
        embed.add_field(
            name="Steps",
            value=f"• Extracting ZIP: {'✅' if current_step > 0 else '⏳'}\n"
                  f"• Parsing Questions: {'✅' if current_step > 1 else '⏳'}\n"
                  f"• Generating Document: {'✅' if current_step > 2 else '⏳'}\n"
                  f"• Adding Code Sections: {'✅' if current_step > 3 else '⏳'}\n"
                  f"• Adding Screenshots: {'✅' if current_step > 4 else '⏳'}\n"
                  f"• Finalizing Report: {'✅' if current_step > 5 else '⏳'}",
            inline=False
        )
        
        await message.edit(embed=embed)

    def create_progress_bar(self, percentage):
        """Create a text-based progress bar."""
        filled_length = int(20 * percentage / 100)
        bar = '█' * filled_length + '-' * (20 - filled_length)
        return f'[{bar}]'

    @commands.command(name='banau')
    async def generate_lab_report(self, ctx):
        """Generate a lab report from a ZIP file containing required files."""
        if not ctx.message.attachments:
            await ctx.send("Please upload a ZIP file containing the required files (questions.txt, code files, and screenshots).")
            return

        zip_attachment = ctx.message.attachments[0]
        if not zip_attachment.filename.lower().endswith('.zip'):
            await ctx.send("Please upload a valid ZIP file.")
            return

        # Create initial progress embed
        progress_embed = discord.Embed(
            title="Lab Report Generation Progress",
            description="[--------------------] 0%\n**Current Task:** Preparing to start",
            color=discord.Color.blue()
        )
        loading_message = await ctx.send(embed=progress_embed)

        try:
            # Step 1: Extract ZIP
            await self.update_progress_embed(loading_message, 1, 6, "Extracting ZIP file")
            zip_bytes = await zip_attachment.read()
            extracted_files, temp_dir = self.extract_zip(zip_bytes)

            # Step 2: Parse Questions
            await self.update_progress_embed(loading_message, 2, 6, "Parsing questions file")
            questions_file = None
            code_files = {}
            screenshot_files = {}
            missing_files = []

            for filename, filepath in extracted_files.items():
                if filename == 'questions.txt':
                    questions_file = filepath
                elif any(filename.endswith(ext) for ext in ['.py', '.c', '.cpp', '.java']):
                    number = self.extract_number_from_filename(filename)
                    if number:
                        code_files[number] = filepath
                elif any(filename.endswith(ext) for ext in ['.png', '.jpg', '.jpeg']):
                    number = self.extract_number_from_filename(filename)
                    if number:
                        screenshot_files.setdefault(number, []).append(filepath)

            if not questions_file:
                await loading_message.delete()
                await ctx.send("Missing questions.txt file!")
                return

            # Step 3: Generate Document
            await self.update_progress_embed(loading_message, 3, 6, "Creating document structure")
            with open(questions_file, 'r', encoding='utf-8') as f:
                questions_content = f.read()

            questions = self.parse_questions(questions_content)

            doc = Document()

            # Set font and formatting
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
            style.paragraph_format.line_spacing = 1

            # Add title
            await self.update_progress_embed(loading_message, 4, 6, "Adding document title")
            title = doc.add_heading('Lab Report', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            title_run = title.runs[0]
            title_run.font.name = 'Times New Roman'
            title_run.font.color.rgb = RGBColor(0, 0, 0)

            # Step 4: Add Question, Theory, Code Sections and Screenshots
            for q_num in sorted(questions.keys()):
                await self.update_progress_embed(loading_message, 4, 6, f"Processing Question {q_num}")
                
                # Question Heading
                question_heading = doc.add_heading(f'{q_num}. {questions[q_num]["text"]}', level=1)
                question_heading_run = question_heading.runs[0]
                question_heading_run.font.name = 'Times New Roman'
                question_heading_run.font.color.rgb = RGBColor(0, 0, 0)
                doc.add_paragraph()

                # Theory Section
                if questions[q_num]['theory']:
                    theory_heading = doc.add_heading('Theory:', level=3)
                    theory_heading_run = theory_heading.runs[0]
                    theory_heading_run.font.name = 'Times New Roman'
                    theory_heading_run.font.color.rgb = RGBColor(0, 0, 0)
                    doc.add_paragraph()

                    
                    theory_para = doc.add_paragraph(questions[q_num]['theory'])
                    theory_para.paragraph_format.line_spacing = 1.5
                    doc.add_paragraph()  # Add a blank line after theory

                # Add Code Section
                if q_num in code_files:
                    with open(code_files[q_num], 'r', encoding='utf-8') as code_file:
                        code_content = code_file.read()

                    code_heading = doc.add_heading('Code:', level=3)
                    code_heading_run = code_heading.runs[0]
                    code_heading_run.font.name = 'Times New Roman'
                    code_heading_run.font.color.rgb = RGBColor(0, 0, 0)
                    doc.add_paragraph()

                    code_para = doc.add_paragraph()
                    for line in code_content.splitlines():
                        code_run = code_para.add_run(line + '\n')
                        code_run.font.name = 'Consolas'
                        code_run.font.size = Pt(10)

                    code_para.paragraph_format.line_spacing = 1
                    code_para.paragraph_format.space_after = Pt(0)
                    
                    # Add page break after code section
                    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                else:
                    missing_files.append(f"Code file for Question {q_num}")

                # Add Screenshots for this question
                if q_num in screenshot_files:
                    for i, screenshot_path in enumerate(screenshot_files[q_num]):

                        output_heading = doc.add_heading('Output:', level=3)
                        output_heading_run = output_heading.runs[0]
                        output_heading_run.font.name = 'Times New Roman'
                        output_heading_run.font.color.rgb = RGBColor(0, 0, 0)

                        output_paragraph = doc.add_paragraph()

                        with open(screenshot_path, 'rb') as screenshot_file:
                            screenshot_bytes = screenshot_file.read()
                            doc.add_picture(io.BytesIO(screenshot_bytes), width=Inches(6))

                    # Add page break after screenshots for this question
                    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                else:
                    missing_files.append(f"Screenshot for Question {q_num}")

            # Step 5: Finalize Report
            await self.update_progress_embed(loading_message, 6, 6, "Saving document")
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)

            # Create final embed
            result_embed = discord.Embed(
                title="Lab Report Generated Successfully",
                description="Your lab report has been created!",
                color=discord.Color.green()
            )
            
            if missing_files:
                result_embed.add_field(
                    name="Missing Files", 
                    value="\n".join(missing_files), 
                    inline=False
                )

            result_embed.set_footer(text="File: lab_report.docx")

            # Delete the loading message
            await loading_message.delete()

            await ctx.send(
                embed=result_embed,
                file=discord.File(doc_buffer, filename="lab_report.docx")
            )

        except Exception as e:
            # Create error embed
            error_embed = discord.Embed(
                title="Lab Report Generation Failed",
                description=f"An error occurred: {str(e)}",
                color=discord.Color.red()
            )
            await loading_message.delete()
            await ctx.send(embed=error_embed)

# Add the cog to the bot
async def setup(bot):
    await bot.add_cog(LabReportGenerator(bot))